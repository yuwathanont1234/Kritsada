"""Build the comprehensive market research + financial projection report (.docx)."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from model import (
    CONSERVATIVE, REALISTIC, OPTIMISTIC,
    run_month_by_month,
    TIER_PRICE_THB, TIER_AVG_SCANS, COST_PER_SCAN_THB,
    FIXED_COSTS_THB, APP_STORE_TAKE_RATE,
)

doc = Document()

# ── Document defaults — Thai-friendly font ──────────────────────────
style = doc.styles['Normal']
style.font.name = 'TH Sarabun New'
style.font.size = Pt(14)
rpr = style.element.get_or_add_rPr()
rfonts = OxmlElement('w:rFonts')
rfonts.set(qn('w:ascii'), 'TH Sarabun New')
rfonts.set(qn('w:hAnsi'), 'TH Sarabun New')
rfonts.set(qn('w:cs'), 'TH Sarabun New')
rpr.append(rfonts)

# Margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = 'TH Sarabun New'
        run.font.size = Pt(20 if level == 1 else 16 if level == 2 else 14)
        run.font.color.rgb = RGBColor(0x9A, 0x73, 0x26)  # gold
        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:cs'), 'TH Sarabun New')
        rPr.append(rFonts)


def body(text, bold=False, italic=False, size=14):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'TH Sarabun New'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:cs'), 'TH Sarabun New')
    rPr.append(rFonts)
    return p


def bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    for run in p.runs:
        run.font.name = 'TH Sarabun New'
        run.font.size = Pt(14)
        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:cs'), 'TH Sarabun New')
        rPr.append(rFonts)


def make_table(headers, rows, col_widths_cm=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'

    # Header row
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.name = 'TH Sarabun New'
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Set cell bg color
        tc_pr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1C1610')
        tc_pr.append(shd)
    # Body rows
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ''
            p = cells[c_idx].paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = 'TH Sarabun New'
            run.font.size = Pt(12)
            rPr = run._element.get_or_add_rPr()
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:cs'), 'TH Sarabun New')
            rPr.append(rFonts)
    # Column widths
    if col_widths_cm:
        for row in table.rows:
            for i, w in enumerate(col_widths_cm):
                row.cells[i].width = Cm(w)


# ═══════════════════════════════════════════════════════════════════
# COVER
# ═══════════════════════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Luxury Authenticator")
run.font.size = Pt(32)
run.font.bold = True
run.font.color.rgb = RGBColor(0x9A, 0x73, 0x26)
run.font.name = 'TH Sarabun New'

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("รายงานวิจัยการตลาด + งบประมาณการรายเดือน ปีที่ 1")
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.font.name = 'TH Sarabun New'

doc.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run("Market Research & Year-1 Financial Projection\n"
                   "3 Scenarios: Conservative / Realistic / Optimistic\n\n"
                   "วันที่: 26 พฤษภาคม 2026")
run.font.size = Pt(14)
run.font.italic = True
run.font.name = 'TH Sarabun New'

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════════
heading("1. บทสรุปผู้บริหาร (Executive Summary)", 1)

body("รายงานฉบับนี้ประเมินความเป็นไปได้ของการดำเนินธุรกิจ Luxury Authenticator "
     "ในรูปแบบ Subscription mobile app สำหรับตลาดประเทศไทย โดยใช้ข้อมูลจาก:")
bullet("Statista — Thailand luxury watch market 2025: $310M USD, CAGR 4.01%")
bullet("IMARC — Thailand total watch market: $327M (2025) → $509M (2034)")
bullet("RevenueCat State of Subscription Apps 2025 — SEA freemium conversion 1-3%")
bullet("Entrupy / LegitCheck / Chrono24 — competitor pricing analysis")
bullet("Thai customs / SCMP — counterfeit market intelligence")
bullet("Cost structure ที่เราออกแบบ + optimizations แล้ว (Price cache, Auth bypass, Flash routing)")

body("")
body("ผลการประเมิน 3 Scenarios — กำไร/ขาดทุนสุทธิปีที่ 1:", bold=True)

scenarios_data = []
for s in [CONSERVATIVE, REALISTIC, OPTIMISTIC]:
    rows = run_month_by_month(s)
    total_gross = sum(r['gross_revenue'] for r in rows)
    total_net = sum(r['net_revenue'] for r in rows)
    total_cost = sum(r['total_costs'] for r in rows)
    profit = sum(r['net_profit'] for r in rows)
    scenarios_data.append([
        s.label_th.split(' (')[0],
        f"{rows[-1]['cumulative_installs']:,}",
        f"{rows[-1]['active_paid_total']:,}",
        f"฿{total_gross:,.0f}",
        f"฿{total_cost:,.0f}",
        f"฿{profit:,.0f}",
    ])

make_table(
    ['Scenario', 'Installs Y1', 'Paid (M12)', 'Gross Rev', 'Total Cost', 'Net Profit'],
    scenarios_data,
    col_widths_cm=[3, 2.5, 2.2, 2.8, 2.8, 2.8],
)

body("")
body("ข้อสรุปสำคัญ:", bold=True)
bullet("Conservative: ขาดทุน ~฿30,000 ในปีแรก (ยังอยู่ในระยะ growth — ต้องไป Y2 break-even)")
bullet("Realistic: กำไร ~฿970,000 ในปีแรก — break-even ภายในเดือน M4-M5")
bullet("Optimistic: กำไร ~฿9.7M ในปีแรก — ขึ้นกับ user acquisition + word of mouth")
bullet("ตัวขับกำไรหลัก: tier mix (% Premium users) + churn rate + scan cache hit ratio")
bullet("Cost per scan ที่เราออกแบบ (฿0.90 steady state) ทำให้ unit economics positive ตั้งแต่ tier Standard")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 2. MARKET RESEARCH
# ═══════════════════════════════════════════════════════════════════
heading("2. การวิจัยตลาด (Market Research)", 1)

heading("2.1 ขนาดตลาดประเทศไทย", 2)
body("ตลาดนาฬิกาหรูประเทศไทยเป็นตลาด niche แต่มีกำลังซื้อสูงและเติบโตต่อเนื่อง:")
make_table(
    ['ดัชนี', 'มูลค่า 2025', 'มูลค่าคาดการณ์', 'CAGR', 'แหล่งข้อมูล'],
    [
        ['Total Watch Market TH', '$327M', '$509M (2034)', '4.78%', 'IMARC'],
        ['Luxury Watch Market TH', '$310M', '$378M (2030)', '4.01%', 'Statista'],
        ['Smartphone Users TH (2026)', '60M users', '—', '—', 'Statista'],
        ['Online Watch Buyers', '~5% of luxury buyers', '+50% by 2030', '~8%', 'Estimate'],
    ],
    col_widths_cm=[5, 3, 3.5, 1.8, 3],
)

heading("2.2 ปัญหาตลาด — สินค้าปลอมระบาด", 2)
body("ประเทศไทยเป็นที่รู้จักในระดับโลกในฐานะแหล่งซื้อนาฬิกาปลอมราคาถูก:")
bullet("MBK Center (กรุงเทพ), Chinatown, Bang Niang (Phuket) เป็นจุดขายหลัก")
bullet("กระทรวงการคลังของ Thai Customs ยอมรับเคยขาย Rolex / Patek / Richard Mille ปลอมในการประมูล (2025)")
bullet("ปฏิบัติการล่าสุดยึดนาฬิกาปลอม 809 เรือนใน Chinatown raid ตุลาคม 2025")
bullet("ระดับโลก: 50% ของของปลอมทั้งหมดคือ Rolex (รายงาน SCMP)")
bullet("ผู้เชี่ยวชาญแยกของปลอมด้วยตาเปล่าได้จาก 80% เหลือเพียง 20% (super clone กำลังพัฒนา)")

body("")
body("สิ่งนี้สร้าง 'pain point' จริงสำหรับผู้ซื้อรายย่อย ผู้สะสมมือใหม่ ดีลเลอร์/พ่อค้ามือสอง "
     "ที่ต้องการเครื่องมือ pre-purchase verification อย่างรวดเร็วและถูก")

heading("2.3 คู่แข่งและตำแหน่งทางการตลาด", 2)
make_table(
    ['คู่แข่ง', 'รูปแบบ', 'ราคา', 'เป้าหมาย', 'ช่องว่างของเรา'],
    [
        ['Entrupy', 'B2B Hardware+SaaS', '$139-$1,049/mo (25-250 tokens)', 'Dealer/Pawn', 'แพง 10x, ฮาร์ดแวร์เฉพาะ'],
        ['LegitCheck (App)', 'Pay-per-auth + sub', '$9.99/mo + $10-20/item', 'Consumer', 'ไม่ realtime, ต้องรอคน'],
        ['LegitGrails', 'Per-item service', '$50-$105/item', 'Premium', 'แพง, ไม่มีแอป'],
        ['eBay Auth Guarantee', 'Embedded', 'Free $2k+ / $80 below', 'eBay buyers', 'จำกัด platform'],
        ['Chrono24 Certified', 'One-off cert', '$249/cert', 'Chrono24', 'แพง, ส่งของจริง'],
        ['Luxury Authenticator', 'AI subscription', '฿990-4,990/mo', 'TH market', 'instant + ถูกกว่า 5-10x'],
    ],
    col_widths_cm=[2.5, 2.5, 3.5, 2.5, 4],
)

body("")
body("จุดเด่นของเราในตลาด:", bold=True)
bullet("Real-time AI authentication (ไม่ต้องรอคน 24-48 ชม. แบบ LegitCheck)")
bullet("ราคาถูกกว่า Entrupy 10 เท่า — เข้าถึงผู้ใช้รายบุคคล ไม่จำกัดเฉพาะ dealer")
bullet("ภาษาไทย + UI สำหรับนักสะสมไทยโดยเฉพาะ")
bullet("Heatmap overlay + cost-per-scan ต่ำมาก (฿0.30-2.98) เปรียบเทียบ $10/scan ของคู่แข่ง")

heading("2.4 ตลาดเป้าหมายและ Persona", 2)
body("Total Addressable Market (TAM) ในประเทศไทย — แบ่งตาม persona:")
make_table(
    ['Persona', 'จำนวนประมาณ', 'พฤติกรรม', 'แพ็คเกจที่เหมาะ'],
    [
        ['Casual Collector\n(1-3 watches)', '~150,000 คน', '2-5 scans/mo, verify ก่อนซื้อ', 'Standard ฿990'],
        ['Active Enthusiast\n(4-10 watches)', '~25,000 คน', '10-30 scans/mo, ซื้อ-ขาย', 'Standard / Pro'],
        ['Dealer / Reseller\n(IG sellers, ปั๊มมือสอง)', '~3,000-5,000 ราย', '50-200 scans/mo, listing verify', 'Pro / Premium'],
        ['Watch Boutique\n(ร้านนาฬิกา)', '~300-500 ร้าน', '200+ scans/mo, auth สูง', 'Premium / Enterprise'],
    ],
    col_widths_cm=[3.5, 3, 5, 3.5],
)

body("")
body("Serviceable Addressable Market (SAM) ที่เข้าถึงได้จริงปีแรก: ~5,000-10,000 active users "
     "(0.5-1% ของ TAM ขั้นต่ำ)")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 3. PRICING & SUBSCRIPTION STRATEGY
# ═══════════════════════════════════════════════════════════════════
heading("3. กลยุทธ์การกำหนดราคา (Pricing Strategy)", 1)

heading("3.1 โครงสร้างแพ็คเกจ (ที่ตั้งใน app แล้ว)", 2)
make_table(
    ['Tier', 'ราคา/เดือน', 'Scans/mo', 'Auth/mo', 'ฟีเจอร์เด่น'],
    [
        ['Free',     '฿0',     '5 (per 30d)', '3', 'จำกัด — Bait tier'],
        ['Standard', '฿990',   '50',          '50',  'Pre-purchase verify'],
        ['Pro',      '฿1,990', '100',         '100', '+ Price live, BG removal'],
        ['Premium',  '฿4,990', '200',         '200', '+ Heatmap, PDF, no watermark, cloud backup'],
    ],
    col_widths_cm=[2, 2, 2, 1.8, 6],
)

heading("3.2 ต้นทุนต่อ Scan (Cost per Scan)", 2)
body("หลังจาก optimization ทั้งหมด (Price cache 30d, Cheap-brand bypass, "
     "Pro→Flash routing, Probe gate, MoonSwatch fast-path):")
make_table(
    ['สถานการณ์', 'สัดส่วน', 'ต้นทุน/scan', 'หมายเหตุ'],
    [
        ['Cache hit (scan ซ้ำ)',      '60%',     '฿0.30', 'price + identify cached'],
        ['Cheap-brand bypass',         '25%',     '฿1.10', 'auth ข้าม, identify+price ปกติ'],
        ['Full pipeline',              '15%',     '฿2.98', 'identify+auth+price+grounding'],
        ['Blended (steady state)',     '100%',    f'฿{COST_PER_SCAN_THB:.2f}', '— ใช้ในโมเดล'],
    ],
    col_widths_cm=[5, 2, 2.5, 5],
)

heading("3.3 Unit Economics ต่อ Tier", 2)
make_table(
    ['Tier', 'Price', 'Scan cost', 'Net (หลังหัก App Store)', 'Net Margin', 'LTV (10mo)'],
    [
        ['Standard', '฿990',
         f'฿{TIER_AVG_SCANS["standard"] * COST_PER_SCAN_THB:.0f}',
         f'฿{990 * (1 - APP_STORE_TAKE_RATE):.0f}',
         f'฿{990 * (1 - APP_STORE_TAKE_RATE) - TIER_AVG_SCANS["standard"] * COST_PER_SCAN_THB:.0f}',
         f'฿{(990 * (1 - APP_STORE_TAKE_RATE) - TIER_AVG_SCANS["standard"] * COST_PER_SCAN_THB) * 10:,.0f}'],
        ['Pro', '฿1,990',
         f'฿{TIER_AVG_SCANS["pro"] * COST_PER_SCAN_THB:.0f}',
         f'฿{1990 * (1 - APP_STORE_TAKE_RATE):.0f}',
         f'฿{1990 * (1 - APP_STORE_TAKE_RATE) - TIER_AVG_SCANS["pro"] * COST_PER_SCAN_THB:.0f}',
         f'฿{(1990 * (1 - APP_STORE_TAKE_RATE) - TIER_AVG_SCANS["pro"] * COST_PER_SCAN_THB) * 10:,.0f}'],
        ['Premium', '฿4,990',
         f'฿{TIER_AVG_SCANS["premium"] * COST_PER_SCAN_THB:.0f}',
         f'฿{4990 * (1 - APP_STORE_TAKE_RATE):.0f}',
         f'฿{4990 * (1 - APP_STORE_TAKE_RATE) - TIER_AVG_SCANS["premium"] * COST_PER_SCAN_THB:.0f}',
         f'฿{(4990 * (1 - APP_STORE_TAKE_RATE) - TIER_AVG_SCANS["premium"] * COST_PER_SCAN_THB) * 10:,.0f}'],
    ],
    col_widths_cm=[2, 2, 2.5, 4, 2.5, 3],
)

body("")
body("Net margin ทั้ง 3 tier เป็นบวก — แสดงว่า unit economics ไม่ขึ้นกับ scale "
     "(ยิ่งมี user ยิ่งทำกำไร, ไม่ใช่ขาดทุน per-user)", italic=True)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 4. USER ACQUISITION & FUNNEL ASSUMPTIONS
# ═══════════════════════════════════════════════════════════════════
heading("4. การคาดการณ์การได้ User และ Conversion Funnel", 1)

heading("4.1 Funnel Assumptions", 2)
make_table(
    ['ขั้น', 'Conservative', 'Realistic', 'Optimistic', 'อ้างอิง Benchmark'],
    [
        ['Install → Trial start',   '18%', '25%', '32%', 'Adapty: 10-32% เฉลี่ย'],
        ['Trial → Paid',            '15%', '20%', '27%', 'SEA: 19-26%'],
        ['Install → Paid (net)',    '2.7%', '5.0%', '8.6%', 'SEA freemium 1-5%'],
        ['Monthly churn (paid)',    '10-18%', '8-15%', '6-12%', 'SaaS SEA: 8-15%'],
    ],
    col_widths_cm=[5, 2, 2, 2, 5],
)

heading("4.2 Tier Mix (สัดส่วน paid users)", 2)
make_table(
    ['Tier', 'Conservative', 'Realistic', 'Optimistic', 'หมายเหตุ'],
    [
        ['Standard ฿990',  '75%', '70%', '62%', 'Mass market, ราคาเข้าถึงง่าย'],
        ['Pro ฿1,990',     '20%', '22%', '27%', 'Active collectors / small dealers'],
        ['Premium ฿4,990', '5%',  '8%',  '11%', 'Boutique / heavy dealers'],
    ],
    col_widths_cm=[3, 2, 2, 2, 6],
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 5. MONTHLY P&L (Realistic scenario)
# ═══════════════════════════════════════════════════════════════════
heading("5. งบกำไรขาดทุนรายเดือน — Realistic Scenario", 1)
body("ตารางต่อไปนี้แสดง P&L 12 เดือนสำหรับ scenario กลาง (Base case) "
     "— สำหรับ Conservative และ Optimistic ดูในไฟล์ Excel แนบ", italic=True)

rows = run_month_by_month(REALISTIC)

# Splitting into 2 tables (M1-M6, M7-M12) for page fit
def make_pnl_table(months_subset, m_indices):
    table_rows = [
        ['New installs']           + [f'{r["new_installs"]:,}' for r in months_subset],
        ['Cumulative installs']    + [f'{r["cumulative_installs"]:,}' for r in months_subset],
        ['New paid']               + [f'{r["new_paid"]:,}' for r in months_subset],
        ['Active paid total']      + [f'{r["active_paid_total"]:,}' for r in months_subset],
        [' • Standard']            + [f'{r["active_paid_standard"]:,}' for r in months_subset],
        [' • Pro']                 + [f'{r["active_paid_pro"]:,}' for r in months_subset],
        [' • Premium']             + [f'{r["active_paid_premium"]:,}' for r in months_subset],
        ['Gross revenue']          + [f'฿{r["gross_revenue"]:,}' for r in months_subset],
        ['  - App Store fee 15%']  + [f'฿{r["app_store_fee"]:,}' for r in months_subset],
        ['NET REVENUE']            + [f'฿{r["net_revenue"]:,}' for r in months_subset],
        ['Variable costs (scans)'] + [f'฿{r["variable_costs"]:,}' for r in months_subset],
        ['Fixed costs (infra)']    + [f'฿{r["fixed_costs"]:,}' for r in months_subset],
        ['Marketing']              + [f'฿{r["marketing"]:,}' for r in months_subset],
        ['TOTAL COSTS']            + [f'฿{r["total_costs"]:,}' for r in months_subset],
        ['NET PROFIT / (LOSS)']    + [f'฿{r["net_profit"]:,}' for r in months_subset],
    ]
    headers = ['Metric'] + [f'M{i}' for i in m_indices]
    make_table(headers, table_rows, col_widths_cm=[3.5] + [1.8]*6)

heading("ครึ่งปีแรก (M1-M6)", 3)
make_pnl_table(rows[:6], range(1, 7))
body("")
heading("ครึ่งปีหลัง (M7-M12)", 3)
make_pnl_table(rows[6:], range(7, 13))

body("")
total_gross = sum(r['gross_revenue'] for r in rows)
total_net_rev = sum(r['net_revenue'] for r in rows)
total_var = sum(r['variable_costs'] for r in rows)
total_fixed = sum(r['fixed_costs'] for r in rows)
total_mkt = sum(r['marketing'] for r in rows)
total_profit = sum(r['net_profit'] for r in rows)

body("สรุป Realistic ปีที่ 1:", bold=True, size=15)
bullet(f"Gross Revenue:    ฿{total_gross:,.0f}")
bullet(f"App Store Fee:    -฿{total_gross * APP_STORE_TAKE_RATE:,.0f} (15%)")
bullet(f"Net Revenue:      ฿{total_net_rev:,.0f}")
bullet(f"Variable Costs:   ฿{total_var:,.0f}")
bullet(f"Fixed Costs:      ฿{total_fixed:,.0f}")
bullet(f"Marketing:        ฿{total_mkt:,.0f}")
bullet(f"NET PROFIT:       ฿{total_profit:,.0f}")

# Break-even month
breakeven_month = next((r['month'] for r in rows if r['net_profit'] > 0), None)
body("")
if breakeven_month:
    body(f"Break-even เดือนที่: M{breakeven_month}", bold=True)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 6. RISKS & SENSITIVITY
# ═══════════════════════════════════════════════════════════════════
heading("6. ความเสี่ยงและการวิเคราะห์ Sensitivity", 1)

heading("6.1 ความเสี่ยงสูง", 2)
bullet("Replicate / Gemini ราคาขึ้น — กระทบ cost per scan ทันที, "
       "mitigate ด้วย long-term contract หรือ self-host model")
bullet("คู่แข่งใหญ่ (Entrupy / LegitCheck) เปิดตลาดไทย — ต้อง defend ด้วย UX และ price")
bullet("Apple/Google policy change — IAP commission อาจกลับเป็น 30%")
bullet("Counterfeit market เปลี่ยน method — super clones กำลังท้าทาย AI detection")

heading("6.2 ความเสี่ยงปานกลาง", 2)
bullet("Conversion ต่ำกว่าคาด — ต้องปรับ Free tier ให้ดึงดูดมากขึ้นหรือลด pricing")
bullet("Churn สูง — ต้องเพิ่ม retention features (collection vault, social, sharing)")
bullet("App Store Review ปฏิเสธ — IAP, privacy URL, trademark issues")

heading("6.3 Sensitivity Analysis (Realistic ± 20%)", 2)
make_table(
    ['Variable', '-20%', 'Base', '+20%', 'Impact'],
    [
        ['Cost per scan', '฿0.72', '฿0.90', '฿1.08', 'Profit ±฿20-30k'],
        ['Trial→Paid %',  '16%',   '20%',   '24%',   'Profit ±฿250-400k'],
        ['Tier mix Premium', '6%', '8%',    '10%',   'Profit ±฿150-200k'],
        ['Monthly churn', '6-12%', '8-15%', '10-18%', 'Profit ±฿180-300k'],
        ['Install growth', '-20%', 'base',  '+20%',  'Profit ±฿180-220k'],
    ],
    col_widths_cm=[3.5, 2, 2, 2, 4],
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 7. RECOMMENDATIONS
# ═══════════════════════════════════════════════════════════════════
heading("7. ข้อเสนอแนะเชิงกลยุทธ์", 1)

heading("7.1 ระยะสั้น (M1-M3)", 2)
bullet("เปิด Soft launch กับ watch community (Facebook groups เช่น Watch Society Thailand, "
       "Bangkok Watch Club, Rolex Lovers TH) — รับ feedback + early reviews")
bullet("ทำ 7-day free trial เต็มฟีเจอร์ Premium → conversion rate สูงกว่า Free tier 5 scans")
bullet("Content marketing: เขียน guide 'วิธีสังเกต Rolex ปลอม', 'Top 10 รุ่นที่โดน fake บ่อย'")
bullet("Partner กับ IG resellers ดังๆ ให้ใช้ Premium ฟรี แลกกับการ tag/recommend")

heading("7.2 ระยะกลาง (M4-M9)", 2)
bullet("เปิด Annual subscription discount 20% (฿9,990/Standard, ฿19,990/Pro)")
bullet("Refer-a-friend: ผู้แนะนำได้ 1 เดือนฟรี, เพื่อนใหม่ได้ 50% off เดือนแรก")
bullet("เพิ่ม PDF Certificate (ลูกค้าสามารถส่งให้ผู้ซื้อตอนขายต่อ)")
bullet("เพิ่ม Watch Magazine in-app — drive engagement + DAU")

heading("7.3 ระยะยาว (M10-M12+)", 2)
bullet("ขยาย B2B: Dealer Dashboard plan ฿9,990/mo (500 scans, multi-user, branded PDF)")
bullet("API access สำหรับ Marketplace (Lazada, Shopee Mall) ที่ต้องการ verify ก่อน list")
bullet("ขยายไปกระเป๋า / sneakers (ใช้ DINOv3 weights เดียวกัน, train head ใหม่)")
bullet("ขยาย geographic: SG, MY (luxury market similar, conversion rates สูงกว่า)")

heading("7.4 KPI ที่ต้อง track", 2)
bullet("MRR (Monthly Recurring Revenue) — ดูเทรนด์โต")
bullet("Churn rate by tier — Standard ถ้า churn > 15% = pricing/value mismatch")
bullet("CAC (Customer Acquisition Cost) — ต้อง < LTV/3")
bullet("Cache hit rate — ต้องโต > 60% เพื่อ unit economics ดี")
bullet("Scan-to-save ratio — ผู้ใช้ที่ scan แล้ว save ในตู้สะสม = engaged user")

# ═══════════════════════════════════════════════════════════════════
# APPENDIX
# ═══════════════════════════════════════════════════════════════════
doc.add_page_break()
heading("ภาคผนวก: แหล่งข้อมูล (Sources)", 1)
sources = [
    "Statista — Luxury Watches Thailand Market Forecast 2025-2030",
    "IMARC Group — Thailand Watch Market Size 2025-2034",
    "RevenueCat — State of Subscription Apps 2025",
    "Adapty — Free Trial to Paid Conversion Rates 2026",
    "Business of Apps — App Subscription Trial Benchmarks 2026",
    "Entrupy.com — Pricing & Plans (2025)",
    "LegitCheck.app — Club Pricing & Authentication Services",
    "Chrono24 — Certified Authentication Program ($249/cert)",
    "eBay — Authenticity Guarantee for Watches",
    "South China Morning Post — Counterfeit Luxury Watch Market",
    "Thai Newsroom — Chinatown Counterfeit Watch Raid (Oct 2025)",
    "Time and Tide Watches — Thai Customs Counterfeit Auction Scandal",
    "Apple Developer — Small Business Program (15% commission)",
    "Google Play Developer Policy — 15% Service Fee Tier",
]
for s in sources:
    bullet(s)

# Save
output = '/Users/kritsada/Desktop/Luxury-authenticator/financial_model/Luxury_Authenticator_Year1_Report.docx'
doc.save(output)
print(f"✅ Saved: {output}")
import os
print(f"   Size: {os.path.getsize(output)/1024:.1f} KB")
